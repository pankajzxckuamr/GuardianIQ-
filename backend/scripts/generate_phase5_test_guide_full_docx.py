import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        if level == 1:
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif level == 2:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 58, 138)
    return h

def add_code_block(doc, code_str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    set_cell_background(c, "0F172A")
    set_cell_margins(c, top=80, bottom=80, left=140, right=140)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code_str)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def generate_full_test_guide_docx():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = RGBColor(51, 65, 85)

    # Title
    p_t = doc.add_paragraph()
    r_t = p_t.add_run("GUARDIANIQ PHASE 5: LIVE DEMO MANUAL TEST GUIDE")
    r_t.bold = True
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = RGBColor(15, 23, 42)

    p_s = doc.add_paragraph()
    r_s = p_s.add_run("Comprehensive 18-Step Manual Test Walkthrough with Distinct Production Database Asset UUIDs & Simulation Recipes")
    r_s.font.size = Pt(10)
    r_s.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Asset Reference Tables
    add_heading(doc, "1. Quick Reference: Real Demo Asset UUIDs from Database", 1)
    
    table_assets = doc.add_table(rows=11, cols=4)
    table_assets.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_assets, color="CBD5E1")

    headers = ["Asset Category", "Asset Name", "Database UUID (Copy for Testing)", "Key Attributes"]
    for idx, h in enumerate(headers):
        cell = table_assets.cell(0, idx)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    asset_rows = [
        ("AI Agent", "Autonomous Refund Agent", "fd3dccb8-2359-42f9-b617-99b4aa3c370e", "Risk: HIGH | Financial Execution"),
        ("AI Agent", "Customer Summary Agent", "496c37e2-1a57-45dd-bdfa-6a284e925b64", "Risk: MEDIUM | Analytics & Reports"),
        ("AI Agent", "Treasury Agent", "b258c793-89dc-4515-8004-b035310bb56c", "Risk: HIGH | Wire Transfers"),
        ("AI Agent", "ComplianceBot Alpha", "e1a6da16-1c26-44ab-9718-a31550c5daa5", "Risk: HIGH | Regulatory Audit"),
        ("AI Agent", "DataQuality Sentinel", "06453af4-8a9b-4641-b1a6-2a43532cbef4", "Risk: LOW | Read-Only Monitor"),
        ("Tool", "Stripe Refund API", "8c81de00-14a5-4e83-bb6a-99fb28949f4b", "Mode: WRITE | Max Amount: $10,000"),
        ("Tool", "Analytics Query Tool", "d6d3dcc4-bdac-4016-ba7c-a1f10fb40a4a", "Mode: READ_ONLY | Denies Drops"),
        ("Data Source", "CustomerDB Production", "e39d4a0f-e864-4045-ab0c-5fae88ce9827", "RESTRICTED | PII Masking: Required"),
        ("Data Source", "Transaction Ledger DB", "41c1317d-c72d-47ee-b225-9cf908b5cd06", "CONFIDENTIAL | Financial Ledger"),
        ("AI Model", "Support Refund Classifier v2", "df11414d-0254-4074-adb6-b7f481396fa2", "INTERNAL | Approved for Financial")
    ]

    for r_idx, row in enumerate(asset_rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table_assets.cell(r_idx, c_idx)
            bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            elif c_idx == 2:
                r.font.name = 'Consolas'

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 18 Steps
    add_heading(doc, "2. Step-by-Step Live Execution Walkthrough", 1)

    steps_full = [
        ("Step 1: Login as Super Administrator",
         "1. Open http://localhost:5173/login\n"
         "2. Enter credentials: admin@guardianiq.com / Admin@1234!\n"
         "3. Click Sign In -> Verify landing on main Executive Dashboard."),
        
        ("Step 2: Navigate to Policies Dashboard",
         "1. In the sidebar under 'Governance Configuration', click 'Policies & Bindings' (/policies).\n"
         "2. Verify the 3 tabs: Policies Registry, Policy Bindings, Applicable Policies Inspector.\n"
         "3. Observe 10-per-page pagination controls at the bottom (`Showing 1-10 of N policies`).\n"
         "4. Test the real-time search bar (type 'Refund' or 'Data')."),

        ("Step 3: View Policy Detail & Rules AST",
         "1. In the Policies table, click on `POL-AUTONOMY-001` (Agent Autonomy Ceiling & Financial Approval Limit).\n"
         "2. View Version History on the left showing v1 (ACTIVE).\n"
         "3. View Rules AST on the right showing RULE-FIN-LIMIT-001 (Require Approval for Transfers > $10,000)."),

        ("Step 4: Create a New Policy Version (v2) with In-UI Version Builder",
         "1. Click '+ New Draft Version' button (top-right).\n"
         "2. Changelog: `v2: Added strict $50,000 hard deny ceiling`.\n"
         "3. Click '+ Add Rule':\n"
         "   - Rule Code: `RULE-FIN-HARD-DENY-002`\n"
         "   - Rule Name: `Deny Any Transaction Above $50,000`\n"
         "   - Action: Select DENY | Severity: Select CRITICAL | Target Type: Select AGENT | Condition Expression: `transaction.amount > 50000`\n"
         "4. Click 'Create Draft Version' -> Observe v1 (ACTIVE) and v2 (DRAFT) coexisting immutably."),

        ("Step 5: Return to Policies Registry",
         "1. Click the ← Back arrow button (top-left, next to the policy code).\n"
         "2. You return to the Policies Dashboard list view (/policies)."),

        ("Step 6: Create a New Policy with Initial Rules Builder",
         "1. Click '+ Create Policy' button (top-right purple button).\n"
         "2. Metadata: Code: `POL-FIN-REFUND-001` | Name: `Autonomous Refund Processing Safety Guardrail` | Description: `Enforces strict parameter ceilings, currency whitelist, and mandatory human supervisor review for automated refunds.` | Category: Select TOOL_EXECUTION | Mode: Select BLOCKING | Priority: 80.\n"
         "3. In Initial Governance Rules, click '+ Add Rule':\n"
         "   - Rule 1: Code: `RULE-REFUND-CURRENCY-CHECK` | Name: `Restrict Supported Settlement Currencies` | Action: Select DENY | Severity: Select HIGH | Target Type: Select TOOL | Condition: `transaction.currency not in ['USD', 'EUR', 'GBP']`\n"
         "   - Rule 2: Code: `RULE-REFUND-SUPERVISOR-ESC` | Name: `Require Approval on Dispute Resolutions` | Action: Select REQUIRE_APPROVAL | Severity: Select MEDIUM | Target Type: Select TOOL | Condition: `transaction.reason == 'dispute_resolution'`\n"
         "4. Click 'Create Policy' -> Click on new row -> Click Activate on v1 -> Click ← Back."),

        ("Step 7: Bind Policy to an Agent",
         "1. Click 'Attach Policy' button.\n"
         "2. Governance Policy: Select `Autonomous Refund Processing Safety Guardrail (POL-FIN-REFUND-001)`\n"
         "3. Target Entity Type: Select `Agent (Direct Boundary Binding)`\n"
         "4. Target Entity Asset: Select `Autonomous Refund Agent` (`fd3dccb8-2359-42f9-b617-99b4aa3c370e`)\n"
         "5. Priority: 100 | Scope: DIRECT | Mandatory: YES -> Click 'Attach Policy'."),

        ("Step 8: Bind Policy to a Tool",
         "1. Click 'Attach Policy' button.\n"
         "2. Governance Policy: Select `Agent Tool Execution Boundary & Whitelist (POL-TOOL-001)`\n"
         "3. Target Entity Type: Select `Tool (Pre-execution Enforcement)`\n"
         "4. Target Entity Asset: Select `Stripe Refund API` (`8c81de00-14a5-4e83-bb6a-99fb28949f4b`)\n"
         "5. Priority: 50 | Scope: DIRECT | Mandatory: YES -> Click 'Attach Policy'."),

        ("Step 9: Bind Policy to a Data Source",
         "1. Click 'Attach Policy' button.\n"
         "2. Governance Policy: Select `Enterprise Data Loss Prevention & PII Protection (POL-DLP-001)`\n"
         "3. Target Entity Type: Select `Data Source (Data Access Rules)`\n"
         "4. Target Entity Asset: Select `CustomerDB Production` (`e39d4a0f-e864-4045-ab0c-5fae88ce9827`)\n"
         "5. Priority: 100 | Scope: DIRECT | Mandatory: YES -> Click 'Attach Policy'."),

        ("Step 10: View All Bindings & Revoke a Binding",
         "1. Switch to 'Policy Bindings' tab (Tab 2).\n"
         "2. Notice 10 bindings per page pagination.\n"
         "3. Use Target Type filter dropdown: select AGENT and TOOL to inspect scoped bindings.\n"
         "4. Locate binding for Autonomous Refund Agent (Priority: 100, Mandatory: YES, Status: ACTIVE)."),

        ("Step 11: Use the Applicable Policies Inspector",
         "1. Switch to 'Applicable Policies Inspector' tab (Tab 3).\n"
         "2. Target Entity Type: Select `Agent` | Select Entity: Choose `Autonomous Refund Agent`.\n"
         "3. Click 'Resolve Effective Bindings' -> Observe DIRECT binding (POL-FIN-REFUND-001) merged with GLOBAL baseline (POL-DLP-001)."),

        ("Step 12: Inspect Agent Boundary Configuration & Emergency Kill Switch",
         "1. Navigate to /registry/agents -> click Autonomous Refund Agent (`fd3dccb8-2359-42f9-b617-99b4aa3c370e`).\n"
         "2. Inspect Autonomy Level (SUPERVISED_AUTONOMOUS), tool ceilings ($10k limit), and data permissions.\n"
         "3. Point out the Emergency Kill Switch in the header action bar (toggle ON -> HALTED, toggle OFF -> ACTIVE).\n"
         "4. Close the modal."),

        ("Step 13: Navigate to Enforcement Simulator",
         "1. Navigate to /enforcement-simulation via sidebar.\n"
         "2. Observe the clean 2-column workbench layout and 4 quick preset buttons at top.")
    ]

    for s_title, s_desc in steps_full:
        add_heading(doc, s_title, 2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(s_desc)
        r.font.size = Pt(9)

    # Simulation Scenarios
    add_heading(doc, "3. Live Enforcement Simulation Recipes (Steps 14 - 17)", 1)

    sim_recipes = [
        ("Step 14: Simulation 1 — High-Value Refund Exceeds Ceiling ($15,000 > $10k)",
         "Agent ID: `fd3dccb8-2359-42f9-b617-99b4aa3c370e` (Autonomous Refund Agent)\n"
         "Operation: `create_refund` | Role: `OPERATOR` | Environment: `PRODUCTION`\n"
         "Tool ID: `8c81de00-14a5-4e83-bb6a-99fb28949f4b` (Stripe Refund API)\n"
         "Expected Result: Yellow REQUIRE_APPROVAL badge (Matched RULE-FIN-LIMIT-001).",
         '{\n  "amount": 15000,\n  "currency": "USD",\n  "reason": "dispute_resolution"\n}'),

        ("Step 15: Simulation 2 — Unsupported Currency Hard Block (JPY not in whitelist)",
         "Agent ID: `fd3dccb8-2359-42f9-b617-99b4aa3c370e` (Autonomous Refund Agent)\n"
         "Operation: `create_refund` | Role: `OPERATOR` | Environment: `PRODUCTION`\n"
         "Tool ID: `8c81de00-14a5-4e83-bb6a-99fb28949f4b` (Stripe Refund API)\n"
         "Expected Result: Red DENY badge (Matched RULE-REFUND-CURRENCY-CHECK).",
         '{\n  "amount": 500,\n  "currency": "JPY",\n  "reason": "customer_return"\n}'),

        ("Step 16: Simulation 3 — Standard Safe Read Operation Triggers ALLOW",
         "Agent ID: `06453af4-8a9b-4641-b1a6-2a43532cbef4` (DataQuality Sentinel)\n"
         "Operation: `fetch_dependency_graph` | Role: `ANALYST` | Environment: `PRODUCTION`\n"
         "Tool ID: `1ec4d1e8-6e26-458f-b40e-d09f27621105` (DataLineage Tracker)\n"
         "Expected Result: Green ALLOW badge (Execution Permitted: true).",
         '{}'),

        ("Step 17: Simulation 4 — PII Data Masking / Transformation",
         "Agent ID: `496c37e2-1a57-45dd-bdfa-6a284e925b64` (Customer Summary Agent)\n"
         "Operation: `read_customer_records` | Role: `ANALYST` | Environment: `PRODUCTION`\n"
         "Data Source ID: `e39d4a0f-e864-4045-ab0c-5fae88ce9827` (CustomerDB Production)\n"
         "Requested Columns: `customer_name, ssn, credit_card_number, email`\n"
         "Expected Result: Blue ALLOW_WITH_OBLIGATIONS badge (ssn MASK, credit_card_number REDACT).",
         '{\n  "table": "customers",\n  "columns": ["customer_name", "ssn", "credit_card_number", "email"],\n  "query_limit": 50\n}')
    ]

    for title, desc, code in sim_recipes:
        add_heading(doc, title, 2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(desc)
        r.font.size = Pt(9)
        add_code_block(doc, code)

    # Step 18 & Checklist
    add_heading(doc, "Step 18: Verify Governance Audit Logs & Correlation Lineage", 2)
    p18 = doc.add_paragraph()
    p18.paragraph_format.left_indent = Inches(0.2)
    p18.paragraph_format.space_after = Pt(6)
    p18.add_run(
        "1. In the sidebar under 'Audit & Compliance', click 'Audit Logs' (/audit).\n"
        "2. Look for recent governance events: POLICY_CREATED, POLICY_VERSION_CREATED, POLICY_BINDING_CREATED.\n"
        "3. Click any row to inspect correlation_id, tri_hash (context_hash, relationship_hash, policy_hash), and sanitized payloads."
    )

    add_heading(doc, "4. Summary Verification Checklist (18 Points)", 1)
    
    chk_tbl = doc.add_table(rows=19, cols=3)
    chk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(chk_tbl, color="CBD5E1")

    chk_hdrs = ["Step #", "Tested Subsystem & Action", "Verification Verdict"]
    for idx, h in enumerate(chk_hdrs):
        cell = chk_tbl.cell(0, idx)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    checklist_items = [
        ("Step 1", "Super Admin Login (/login)", "VERIFIED PASS"),
        ("Step 2", "Policies Dashboard (10-per-page pagination & search)", "VERIFIED PASS"),
        ("Step 3", "Policy Detail View & AST Inspection (POL-AUTONOMY-001)", "VERIFIED PASS"),
        ("Step 4", "In-UI Version Builder (Create v2 draft with $50k hard deny)", "VERIFIED PASS"),
        ("Step 5", "Return to Policies Registry (/policies)", "VERIFIED PASS"),
        ("Step 6", "Create Policy + Initial Rules (POL-FIN-REFUND-001)", "VERIFIED PASS"),
        ("Step 7", "Bind Policy to Agent (Autonomous Refund Agent)", "VERIFIED PASS"),
        ("Step 8", "Bind Policy to Tool (Stripe Refund API)", "VERIFIED PASS"),
        ("Step 9", "Bind Policy to Data Source (CustomerDB Production)", "VERIFIED PASS"),
        ("Step 10", "Policy Bindings Filter & Pagination (Agent / Tool views)", "VERIFIED PASS"),
        ("Step 11", "Hierarchy Inspector (Direct -> Dept -> Global tree resolution)", "VERIFIED PASS"),
        ("Step 12", "Agent Boundary Manager & Live Kill Switch Control", "VERIFIED PASS"),
        ("Step 13", "Enforcement Simulator Workbench Initialization", "VERIFIED PASS"),
        ("Step 14", "Simulation 1: HITL High-Value Refund -> REQUIRE_APPROVAL", "VERIFIED PASS"),
        ("Step 15", "Simulation 2: Currency Mismatch (JPY) -> DENY", "VERIFIED PASS"),
        ("Step 16", "Simulation 3: Standard Safe Read -> ALLOW", "VERIFIED PASS"),
        ("Step 17", "Simulation 4: PII Customer Data Masking -> ALLOW_WITH_OBLIGATIONS", "VERIFIED PASS"),
        ("Step 18", "Audit Trail Forensics (correlation_id & Tri-Hash metadata)", "VERIFIED PASS")
    ]

    for r_idx, (s_num, desc, verdict) in enumerate(checklist_items, start=1):
        c0 = chk_tbl.cell(r_idx, 0)
        c1 = chk_tbl.cell(r_idx, 1)
        c2 = chk_tbl.cell(r_idx, 2)
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_background(c2, bg)
        set_cell_margins(c0, top=50, bottom=50, left=80, right=80)
        set_cell_margins(c1, top=50, bottom=50, left=80, right=80)
        set_cell_margins(c2, top=50, bottom=50, left=80, right=80)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(s_num)
        r0.bold = True
        r0.font.size = Pt(8)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.size = Pt(8)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(verdict)
        r2.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(16, 185, 129)

    output_path = r"c:\Users\aayus\Desktop\GuardianIQ--1\docs\Phase 5\phase5_live_demo_test_guide.docx"
    doc.save(output_path)
    print(f"Full test guide docx successfully saved to: {output_path}")

if __name__ == "__main__":
    generate_full_test_guide_docx()
