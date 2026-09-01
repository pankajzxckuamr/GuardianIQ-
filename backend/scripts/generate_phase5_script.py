import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        if level == 1:
            r.font.size = Pt(15)
            r.font.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif level == 2:
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 58, 138)
        elif level == 3:
            r.font.size = Pt(10.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(51, 65, 85)
    return h

def add_speaker_card(doc, slide_num, slide_title, duration, spoken_text, visual_cues, transition_text):
    tbl = doc.add_table(rows=4, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="CBD5E1", sz="6")
    
    # Row 0: Slide Header Banner
    c0 = tbl.cell(0, 0)
    set_cell_background(c0, "1E293B")
    set_cell_margins(c0, top=80, bottom=80, left=150, right=150)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(f"SLIDE {slide_num}: {slide_title.upper()}  |  Target Time: {duration}")
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(255, 255, 255)
    
    # Row 1: Visual & Action Cues
    c1 = tbl.cell(1, 0)
    set_cell_background(c1, "F0F9FF")
    set_cell_margins(c1, top=60, bottom=60, left=150, right=150)
    p1 = c1.paragraphs[0]
    r1_tag = p1.add_run("[VISUAL CUES & ACTIONS] ")
    r1_tag.bold = True
    r1_tag.font.size = Pt(8.5)
    r1_tag.font.color.rgb = RGBColor(2, 132, 199)
    r1_text = p1.add_run(visual_cues)
    r1_text.font.size = Pt(8.5)
    r1_text.font.color.rgb = RGBColor(51, 65, 85)

    # Row 2: Word-for-Word Spoken Script
    c2 = tbl.cell(2, 0)
    set_cell_background(c2, "FFFFFF")
    set_cell_margins(c2, top=100, bottom=100, left=150, right=150)
    p2 = c2.paragraphs[0]
    r2_tag = p2.add_run("[WHAT TO SAY - WORD FOR WORD]\n")
    r2_tag.bold = True
    r2_tag.font.size = Pt(9.5)
    r2_tag.font.color.rgb = RGBColor(15, 23, 42)
    r2_text = p2.add_run(spoken_text)
    r2_text.font.size = Pt(10)
    r2_text.font.color.rgb = RGBColor(30, 41, 59)
    p2.paragraph_format.line_spacing = 1.25

    # Row 3: Transition Cue
    c3 = tbl.cell(3, 0)
    set_cell_background(c3, "F8FAFC")
    set_cell_margins(c3, top=60, bottom=60, left=150, right=150)
    p3 = c3.paragraphs[0]
    r3_tag = p3.add_run("[SMOOTH TRANSITION] ")
    r3_tag.bold = True
    r3_tag.font.size = Pt(8.5)
    r3_tag.font.color.rgb = RGBColor(100, 116, 139)
    r3_text = p3.add_run(transition_text)
    r3_text.font.size = Pt(8.5)
    r3_text.font.italic = True
    r3_text.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def generate_script():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(51, 65, 85)

    # Title Block
    p_title = doc.add_paragraph()
    r_t = p_title.add_run("GUARDIANIQ PHASE 5: PRESENTER SCRIPT & SPEAKER NOTES")
    r_t.bold = True
    r_t.font.size = Pt(18)
    r_t.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    r_s = p_sub.add_run("Slide-by-Slide Word-for-Word Narration, UI Demonstration Guide & Executive Q&A Preparation")
    r_s.font.size = Pt(11)
    r_s.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Presentation Structure
    add_heading(doc, "Presentation Timing & Strategy", 1)
    doc.add_paragraph(
        "• Total Presentation Time: ~15 to 20 Minutes (8 min slides + 7 min live demo + 5 min Q&A).\n"
        "• Tone: Confident, authoritative, senior engineering lead. Focus on safety guarantees, fail-closed security, and zero performance penalty.\n"
        "• Live Demo Hand-off: Transition seamlessly from the architecture slides directly into the live browser at http://localhost:5173."
    )

    add_heading(doc, "Slide-by-Slide Presenter Script", 1)

    slides_data = [
        (1, "Title Slide: GuardianIQ Policy & Runtime ENFORCE Layer", "1.0 min",
         "\"Good morning everyone. Today, I am excited to present Phase 5 of GuardianIQ: our Policy and Runtime Enforcement Layer.\n\n"
         "As enterprise organizations deploy autonomous AI agents with access to real banking APIs, customer databases, and production tools, "
         "the greatest challenge is no longer agent capability—it is deterministic governance and safety. "
         "In this session, we will walk you through the end-to-end architecture we built to ensure that no AI agent can execute an unauthorized action, "
         "leak sensitive PII, or bypass administrative oversight. Let's dive in.\"",
         "Project slide on full screen. Introduce yourself and your engineering team (Aayush, Pankaj, Jitendra).",
         "\"To understand why Phase 5 is critical, let's first look at the core enterprise dilemma in autonomous AI governance.\""),

        (2, "Executive Overview: Solving The AI Governance Gap", "1.5 min",
         "\"Traditional security firewalls inspect static network packets, but AI agents present a fundamentally new threat model. "
         "An agent might generate dynamically constructed tool parameters, attempt to read confidential customer data, or execute high-dollar refunds.\n\n"
         "Without runtime enforcement, companies face three critical risks: first, unbounded autonomy where agents execute actions without human review; "
         "second, privacy exposure where restricted data is passed into external LLMs; and third, Time-of-Check to Time-of-Use vulnerabilities where "
         "approved actions are modified right before execution.\n\n"
         "Phase 5 solves this by introducing a deterministic, fail-closed Policy and Enforcement Gateway. We have achieved 100% test coverage with "
         "82 automated test suites verifying every positive, negative, and security path.\"",
         "Point to the 3 columns: Red (The Challenge), Blue (GuardianIQ Solution), Green (Outcomes).",
         "\"Let's look under the hood at the pipeline that intercepts every governed action.\""),

        (3, "System Topology & Runtime Architecture", "2.0 min",
         "\"This diagram illustrates our end-to-end runtime enforcement pipeline. Notice the fundamental design rule: clients and AI agents can never "
         "call target tools or APIs directly. Every governed request must enter through our server-side Runtime Enforcement Gateway at /api/v1/enforce/execute.\n\n"
         "The pipeline executes in five sequential stages:\n"
         "Stage 1: The request is normalized into a standard Governance Context.\n"
         "Stage 2: Hard Boundary Guards check autonomy limits, active kill switches, tool ceilings, and data classifications.\n"
         "Stage 3: The Policy Binding Resolver traverses our multi-layered policy hierarchy and evaluates rules using a safe AST parser.\n"
         "Stage 4: Our Runtime Authorization Service generates cryptographic SHA-256 hashes and issues a single-use token.\n"
         "Stage 5: Only upon explicit ALLOW does the target adapter execute, followed by immediate transactional outbox event emission.\n\n"
         "Most importantly: if any timeout, syntax error, or network partition occurs, the entire pipeline fails closed to DENY.\"",
         "Trace your laser pointer or cursor across the 5 horizontal stages from Left to Right.",
         "\"Now let's examine how policies are organized and resolved across the enterprise hierarchy.\""),

        (4, "Multi-Layer Policy Hierarchy & Rule Evaluator", "1.5 min",
         "\"In an enterprise, governance policies exist at different levels of the organization. GuardianIQ uses a 3-tier hierarchical resolution order:\n"
         "First, DIRECT bindings attached to a specific Agent or Tool. Second, DEPARTMENT bindings inherited by all department members. And third, GLOBAL bindings that apply across the entire enterprise.\n\n"
         "We implemented a Mandatory Policy Lock: if a parent department or global rule is marked mandatory, lower-level bindings cannot relax or override it.\n\n"
         "When multiple policies apply, our Decision Combiner uses strict mathematical precedence: DENY always wins over REQUIRE_APPROVAL, which wins over ALLOW. "
         "Furthermore, rule evaluation uses safe AST expression parsing—we strictly prohibit Python eval() to eliminate code injection vulnerabilities.\"",
         "Highlight the Direct -> Department -> Global hierarchy card on the left and the AST parser on the right.",
         "\"In addition to declarative policies, we enforce hard security boundaries at the asset level.\""),

        (5, "Hard Security Boundaries: Agent, Tool, Data & Model Guards", "1.5 min",
         "\"Phase 5 introduces four dedicated guardrails that act as non-negotiable security boundaries:\n"
         "1. Agent Autonomy Guard: Enforces autonomy levels. A RECOMMEND_ONLY agent is strictly prevented from executing WRITE calls. If an administrator toggles the one-click Kill Switch, all agent actions halt instantly.\n"
         "2. Tool Permission Guard: Checks tool capability tags and enforces parameter limits—for instance, capping refund API calls at $10,000.\n"
         "3. Data Permission Guard: Enforces data classifications from PUBLIC to RESTRICTED and automatically applies field-level masking or tokenization before exposure.\n"
         "4. Model Provider Guard: Prevents sending proprietary or regulated data to unapproved public LLMs.\"",
         "Point to each of the 4 guard cards: Red (Agent), Blue (Tool), Green (Data), Purple (Model).",
         "\"Next, let's explore our breakthrough security feature: cryptographic TOCTOU protection.\""),

        (6, "Cryptographic TOCTOU & Replay Defense", "1.5 min",
         "\"One of the most insidious vulnerabilities in AI systems is Time-of-Check to Time-of-Use tampering. An agent might seek approval to refund $500, "
         "a manager approves it, but between approval and execution the payload is tampered to $50,000.\n\n"
         "GuardianIQ eliminates this vulnerability with our Multi-Hash Authorization Service. When an approval is created, we generate three distinct SHA-256 hashes: "
         "the context_hash of the exact payload, the relationship_hash of the active graph, and the policy_hash of active rules.\n\n"
         "We issue a single-use Runtime Authorization Token with a 60-second TTL. At execution time, the gateway re-hashes the live request. If a single character has changed, "
         "the hashes mismatch and execution is instantly rejected. Once used, the token is timestamped consumed_at, blocking all replay attacks.\"",
         "Emphasize the Tri-Hash formula and the 60-second single-use token lifecycle.",
         "\"To enable developers and compliance officers to test these rules safely, we built the Enforcement Simulator.\""),

        (7, "Developer & Operator Experience: Enforcement Simulator", "1.0 min",
         "\"Testing complex multi-agent policies should never put live production systems at risk. We created the Enforcement Simulation workbench.\n\n"
         "The simulator provides a non-authoritative execution environment where engineers can select any Agent and Tool, paste a test JSON payload, "
         "and receive a full multi-layer decision trace in under 50 milliseconds.\n\n"
         "It breaks down which boundary checks passed, which AST policy rules triggered, what decision was reached (ALLOW, REQUIRE_APPROVAL, or DENY), "
         "and what masking obligations must be applied—all with guaranteed zero target side effects. We will demonstrate this live in just a moment.\"",
         "Point to the 4 simulator benefit tiles and preview the live simulation screen.",
         "\"Before we jump into the live demo, let's review our automated test evidence and validation scorecard.\""),

        (8, "Verification Evidence: 82/82 Automated Tests Passing", "1.0 min",
         "\"Quality engineering and rigorous verification are at the heart of Phase 5. We built an automated test matrix covering 22 comprehensive QA scenarios "
         "specified in our engineering plan.\n\n"
         "Across our 20 test suites in backend/app/tests, all 82 tests pass cleanly in 6.5 seconds. This includes database migration rollbacks, multi-tenant isolation, "
         "expired binding purges, parameter limit overruns, payload tampering rejections, in-memory cache fallbacks, and fail-closed timeout handlers.\n\n"
         "Phase 5 is robust, hardened, and verified with zero open critical defects.\"",
         "Direct attention to the 8 green/blue metric cards and the 100% Pass badge.",
         "\"Now, let's switch to the live application and see Phase 5 in action.\""),

        (9, "Live Demonstration Roadmap", "0.5 min",
         "\"In our live walkthrough, we will demonstrate 6 core capabilities:\n"
         "1. Policies Registry & 10-per-page pagination with target filtering.\n"
         "2. In-UI Rule Builder & Versioning.\n"
         "3. Attaching policies to real agents and inspecting the hierarchical tree.\n"
         "4. Agent Boundary configurations including the live Kill-Switch toggle.\n"
         "5. Live interactive Enforcement Simulations for Permitted, High-Value Approval, Masked, and Denied requests.\n"
         "6. Audit forensics verifying our unified correlation_id event trail.\n\n"
         "Let's switch over to the browser.\"",
         "Transition to browser window open at http://localhost:5173.",
         "Switching to live UI...")
    ]

    for s_num, s_title, s_dur, spoken, cues, trans in slides_data:
        add_speaker_card(doc, s_num, s_title, s_dur, spoken, cues, trans)

    # Q&A Prep Section
    add_heading(doc, "Anticipated Executive & Technical Q&A Preparation", 1)

    qa_items = [
        ("Q1: What happens if the Policy Engine or Redis cache crashes during a live request?",
         "Answer: GuardianIQ is architected with strict fail-closed defaults. If the cache is unavailable, the gateway automatically falls back to authoritative database queries. If the database or engine times out during a critical WRITE operation, the request is immediately blocked (DENY), preventing any un-governed target execution."),
        ("Q2: Can an agent or frontend developer bypass the Enforcement Gateway and call a tool API directly?",
         "Answer: No. Target execution adapters (like Stripe, Core Banking, or Internal Databases) are secured behind server-side private network boundaries. Only the Runtime Enforcement Gateway holds the credentials and runtime authorization service keys to execute adapter actions."),
        ("Q3: How does GuardianIQ prevent Time-of-Check to Time-of-Use (TOCTOU) tampering?",
         "Answer: When a human approves an action, we seal the approval with a SHA-256 hash of the exact payload. At execution time, the gateway re-hashes the incoming payload. If even a single character or dollar amount was modified, the hash fails, the authorization token is invalidated, and execution is rejected."),
        ("Q4: Does adding multi-layer policy checks introduce noticeable latency to AI agent execution?",
         "Answer: No. Active policy versions and boundary rules are cached in high-speed memory with automatic TTL and event-driven invalidation. Full hierarchical policy resolution and AST evaluation executes in under 15 milliseconds, which is negligible compared to standard LLM inference times.")
    ]

    for q, a in qa_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r_q = p.add_run(q)
        r_q.bold = True
        r_q.font.color.rgb = RGBColor(30, 58, 138)
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.paragraph_format.space_after = Pt(4)
        r_a = p_a.add_run(a)
        r_a.font.size = Pt(9.5)

    output_path = r"c:\Users\aayus\Desktop\GuardianIQ--1\docs\Phase 5\GuardianIQ_Phase5_Presenter_Script.docx"
    doc.save(output_path)
    print(f"Presenter script successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_script()
