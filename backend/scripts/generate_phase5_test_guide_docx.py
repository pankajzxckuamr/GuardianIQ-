import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        if level == 1:
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif level == 2:
            r.font.size = Pt(11.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 58, 138)
    return h

def generate_test_guide_docx():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = RGBColor(51, 65, 85)

    # Title
    p_t = doc.add_paragraph()
    r_t = p_t.add_run("GUARDIANIQ PHASE 5: LIVE DEMO MANUAL TEST GUIDE")
    r_t.bold = True
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = RGBColor(15, 23, 42)

    p_s = doc.add_paragraph()
    r_s = p_s.add_run("Step-by-Step Live Execution Scenarios, Real Database Asset UUIDs & Payload Recipes")
    r_s.font.size = Pt(10)
    r_s.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Asset Reference Tables
    add_heading(doc, "1. Live Demo Real Assets Reference Table", 1)
    
    table_assets = doc.add_table(rows=10, cols=4)
    table_assets.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_assets, color="CBD5E1")

    headers = ["Asset Category", "Asset Name", "Database UUID (Copy-Paste)", "Key Attributes"]
    for idx, h in enumerate(headers):
        cell = table_assets.cell(0, idx)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    asset_rows = [
        ("AI Agent", "Autonomous Refund Agent", "fd3dccb8-2359-42f9-b617-99b4aa3c370e", "Risk: HIGH | Autonomy: WITH_BOUNDS"),
        ("AI Agent", "Customer Summary Agent", "496c37e2-1a57-45dd-bdfa-6a284e925b64", "Risk: MEDIUM | Analytics & Reports"),
        ("AI Agent", "Treasury Agent", "b258c793-89dc-4515-8004-b035310bb56c", "Risk: HIGH | Wire & Settlement"),
        ("AI Agent", "DataQuality Sentinel", "06453af4-8a9b-4641-b1a6-2a43532cbef4", "Risk: LOW | Read-Only Monitor"),
        ("Tool", "Stripe Refund API", "8c81de00-14a5-4e83-bb6a-99fb28949f4b", "Mode: WRITE | Max Amount: $10,000"),
        ("Tool", "Analytics Query Tool", "d6d3dcc4-bdac-4016-ba7c-a1f10fb40a4a", "Mode: READ_ONLY | Denies Drops"),
        ("Data Source", "CustomerDB Production", "e39d4a0f-e864-4045-ab0c-5fae88ce9827", "RESTRICTED | PII Masking: Required"),
        ("Data Source", "Transaction Ledger DB", "41c1317d-c72d-47ee-b225-9cf908b5cd06", "CONFIDENTIAL | Financial Ledger"),
        ("AI Model", "Support Refund Classifier v2", "df11414d-0254-4074-adb6-b7f481396fa2", "INTERNAL | Approved for Financial")
    ]

    for r_idx, row in enumerate(asset_rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table_assets.cell(r_idx, c_idx)
            bg = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(30, 58, 138)
            elif c_idx == 2:
                r.font.name = 'Consolas'

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 8 Steps Walkthrough
    add_heading(doc, "2. Live Demo Step-by-Step Execution Sequence", 1)

    steps = [
        ("Step 1: Administrator Sign In",
         "1. Navigate to http://localhost:5173/login\n"
         "2. Login with admin@guardianiq.com / Admin@1234!\n"
         "3. Verify landing on main Executive Dashboard with system metrics."),
        ("Step 2: Policies & Bindings Explorer (10 per page)",
         "1. Navigate to /policies via sidebar.\n"
         "2. Observe 10 policies per page with pagination controls.\n"
         "3. Switch to 'Policy Bindings' tab; filter by 'AI Agents (AGENT)' and test real-time search."),
        ("Step 3: Policy Versioning & In-UI Rule Builder",
         "1. Click on POL-AUTONOMY-001 in the table.\n"
         "2. Click '+ New Draft Version'; add rule RULE-FIN-HARD-DENY-002 (transaction.amount > 50000 -> DENY).\n"
         "3. Click 'Create Draft Version'; observe v1 (ACTIVE) and v2 (DRAFT) coexisting immutably."),
        ("Step 4: Attach Policy to Agent",
         "1. Click 'Attach Policy' button.\n"
         "2. Target: Autonomous Refund Agent (fd3dccb8-2359-42f9-b617-99b4aa3c370e).\n"
         "3. Policy: POL-AUTONOMY-001, Priority: 100, Mandatory: YES. Click Attach."),
        ("Step 5: Applicable Policies Hierarchy Inspector",
         "1. Click Tab 3 (Applicable Policies Inspector).\n"
         "2. Select Agent -> Autonomous Refund Agent -> Click Resolve Policies Hierarchy.\n"
         "3. Verify DIRECT binding (POL-AUTONOMY-001) merged with GLOBAL baseline (POL-DLP-001)."),
        ("Step 6: Agent Boundary & Live Kill Switch",
         "1. Navigate to /registry/agents -> click Autonomous Refund Agent.\n"
         "2. Review Autonomy Level (AUTONOMOUS_WITH_BOUNDS), tool ceilings ($10k), and data permissions.\n"
         "3. Point out the Emergency Kill Switch in the header action bar."),
        ("Step 7: Interactive Enforcement Simulator (4 Scenarios)",
         "Navigate to /enforcement-simulation and test 4 scenarios:\n"
         "• Scenario A (Permitted): Refund Agent + Stripe API ($2,500) -> ALLOW (Green).\n"
         "• Scenario B (HITL Approval): Refund Agent + Stripe API ($15,000) -> REQUIRE_APPROVAL (Yellow).\n"
         "• Scenario C (PII Masking): Customer Summary Agent + CustomerDB -> ALLOW_WITH_OBLIGATIONS (Blue, SSN Masked).\n"
         "• Scenario D (Mode Violation): DataQuality Sentinel + Analytics Tool (DROP_TABLE) -> DENY (Red)."),
        ("Step 8: Live Audit Forensics",
         "1. Navigate to /audit.\n"
         "2. Click any runtime enforcement log row -> open JSON drawer.\n"
         "3. Point out the correlation_id, tri_hash (context_hash, relationship_hash, policy_hash), and masked secrets.")
    ]

    for s_title, s_desc in steps:
        add_heading(doc, s_title, 2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(s_desc)
        r.font.size = Pt(9)

    output_path = r"c:\Users\aayus\Desktop\GuardianIQ--1\docs\Phase 5\phase5_live_demo_test_guide.docx"
    doc.save(output_path)
    print(f"Test guide docx successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_test_guide_docx()
