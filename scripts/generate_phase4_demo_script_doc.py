"""
Python script to generate a professional Word Document (.docx) demo script for
GuardianIQ Phase 4 — Governance & Event Explorer Live Demonstration.
"""
import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text_p_list, title="DEMO TIP / KEY TALKING POINT", border_hex="0284C7", bg_hex="F0F9FF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Set left border thick
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    
    for idx, t in enumerate(text_p_list):
        if idx > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
        run = p.add_run(t)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def build_demo_doc(output_path):
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("GUARDIANIQ PHASE 4")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Navy #0F172A

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Executive Live Demonstration Script — Event Explorer & AI Governance Ledger")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x02, 0x84, 0xC7) # Sky Blue #0284C7

    # Metadata Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Release Version", "GuardianIQ v4.0.0-rc1 (Phase 4 Audit & Governance Package)"),
        ("Featured AI Agent Entity", "Financial Analysis Bot (ID: d6a3cb9e-11a8-4004-b82f-a38e33790df0)"),
        ("Featured AI Model Entity", "GPT-4o Enterprise Model (ID: mdl_gpt4_gov)"),
        ("Trace Correlation ID", "5c3c5751-3232-4a3f-85ec-247d55077c03"),
        ("Target Audience", "Enterprise Security Auditors, Compliance Officers, PMs, Executive Leadership")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        c0 = meta_table.cell(row_idx, 0)
        c1 = meta_table.cell(row_idx, 1)
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for section headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        return h

    # Section 1: Demonstration Overview & Storyline
    add_heading_1("1. Demonstration Overview & Storyline")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "This demonstration script guides presenters through a end-to-end walkthrough of GuardianIQ Phase 4's "
        "Audit & Governance module. The scenario models an enterprise financial workflow executed by an autonomous AI Agent ("
    )
    r_agent = p.add_run("Financial Analysis Bot")
    r_agent.bold = True
    p.add_run(") connected to an enterprise AI Model (")
    r_model = p.add_run("GPT-4o Enterprise")
    r_model.bold = True
    p.add_run(
        "). Presenters will demonstrate live event stream exploration, cryptographic hash validation, multi-step "
        "timeline tracing, policy violation interception, dead-letter outbox recovery, and automated compliance export generation."
    )

    add_callout(
        doc,
        [
            "Ensure the local stack is running (Backend API on port 8000, React Frontend on port 5173).",
            "Seed database events by executing: python backend/scripts/populate_phase4_events.py before starting the demo."
        ],
        title="PREREQUISITE SETUP"
    )

    # Section 2: Step-by-Step Live Demo Acts
    add_heading_1("2. Step-by-Step Live Demonstration Script")

    # ACT 1
    add_heading_2("Act 1: Event Explorer & Live Stream Filtering (/audit)")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("📍 ").bold = True
    p.add_run("Frontend Route: ").bold = True
    p.add_run("/audit (Page Component: AuditPage.tsx)\n")
    p.add_run("🎯 Objective: ").bold = True
    p.add_run("Show high-throughput searchable governance stream & multi-dimensional filters.")

    # Dialogue Table Act 1
    t1 = doc.add_table(rows=3, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Demonstration Step & Action", "Presenter Script & Spoken Narrative"]
    for c_idx, text in enumerate(headers):
        cell = t1.cell(0, c_idx)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_h = p_h.add_run(text)
        r_h.bold = True
        r_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h.font.size = Pt(10)
    
    act1_steps = [
        (
            "Step 1.1: Open Event Explorer\nNavigating to /audit dashboard.",
            "\"Welcome everyone. What you see on screen is the GuardianIQ Phase 4 Event Explorer. Every interaction across our enterprise AI agents, models, policy checks, and boundary transfers is immutably captured into canonical Event Envelope 2.0 structures in real time.\""
        ),
        (
            "Step 1.2: Demonstrate Live Filtering\nFilter Category to 'Agent' & Classification to 'INTERNAL'.",
            "\"Let's filter down our governance stream. As I select Category: 'Agent' and Classification: 'INTERNAL', you can see the table instantly updates. We can pinpoint specific entity streams, such as our Financial Analysis Bot (ID: d6a3cb9e-11a8-4004-b82f-a38e33790df0).\""
        )
    ]
    for r_idx, (step_txt, script_txt) in enumerate(act1_steps, start=1):
        c0 = t1.cell(r_idx, 0)
        c1 = t1.cell(r_idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(step_txt)
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(script_txt)
        r1.font.size = Pt(9.5)
        r1.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ACT 2
    add_heading_2("Act 2: Canonical Event Envelope Inspection (EventDrawer.tsx)")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("📍 ").bold = True
    p.add_run("Frontend Component: ").bold = True
    p.add_run("EventDrawer.tsx (Triggered from row click)\n")
    p.add_run("🎯 Objective: ").bold = True
    p.add_run("Inspect 20-field canonical event envelope, cryptographic SHA-256 hash, and payload redaction.")

    t2 = doc.add_table(rows=3, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, text in enumerate(headers):
        cell = t2.cell(0, c_idx)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_h = p_h.add_run(text)
        r_h.bold = True
        r_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h.font.size = Pt(10)

    act2_steps = [
        (
            "Step 2.1: Open Slide-over Drawer\nClicking row WORKFLOW_RUN_STARTED for agent d6a3cb9e...",
            "\"By clicking any row, GuardianIQ launches our shared EventDrawer component. Notice the prominent green badge at the top: 'Canonical Envelope 2.0 Hash Verified'. GuardianIQ computes a SHA-256 hash over every event envelope upon ingestion. If any payload or metadata field were tampered with, this hash validation would fail immediately.\""
        ),
        (
            "Step 2.2: Highlight Actor & Subject Context\nExamine Actor JSON and Subject JSON sections.",
            "\"Notice the rich actor context: User ID usr_admin_01 initiated this run with role ADMIN from IP 192.168.1.45. The subject entity is explicitly tagged as entity_type: 'agents' and entity_id: 'd6a3cb9e-11a8-4004-b82f-a38e33790df0'. Below, sensitive payload attributes are automatically masked with [REDACTED] to maintain data privacy compliance.\""
        )
    ]
    for r_idx, (step_txt, script_txt) in enumerate(act2_steps, start=1):
        c0 = t2.cell(r_idx, 0)
        c1 = t2.cell(r_idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(step_txt)
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(script_txt)
        r1.font.size = Pt(9.5)
        r1.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ACT 3
    add_heading_2("Act 3: Subject & Correlation Timeline Visualizer")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("📍 ").bold = True
    p.add_run("Frontend Routes: ").bold = True
    p.add_run("/audit/timeline/agents/d6a3cb9e-11a8-4004-b82f-a38e33790df0 & /audit/events/correlation/5c3c5751-3232-4a3f-85ec-247d55077c03\n")
    p.add_run("🎯 Objective: ").bold = True
    p.add_run("Demonstrate chronological visual lineage & causation chain for the AI Agent.")

    t3 = doc.add_table(rows=3, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, text in enumerate(headers):
        cell = t3.cell(0, c_idx)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_h = p_h.add_run(text)
        r_h.bold = True
        r_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h.font.size = Pt(10)

    act3_steps = [
        (
            "Step 3.1: Navigate to Subject Timeline\nOpening Subject Timeline page for agent d6a3cb9e...",
            "\"Now, let's step into the Subject Timeline view for our Financial Analysis Bot. GuardianIQ reconstructs a vertical chronological graph of every lifecycle event. We see 7 distinct steps: Run Ingestion -> Policy Evaluation (POL-DATA-PRIVACY) -> Model Binding (GPT-4o Enterprise) -> Step Start -> Boundary Check -> Human Approval -> Completion.\""
        ),
        (
            "Step 3.2: Drill into Correlation ID Trace\nClicking 'Trace Correlation ID 5c3c5751...'",
            "\"Notice how every step shares Correlation ID: 5c3c5751-3232-4a3f-85ec-247d55077c03. By clicking Correlation Trace, GuardianIQ instantly isolated the multi-service execution path across workflow_scheduler, policy_engine, relationship_service, and approval_engine within milliseconds.\""
        )
    ]
    for r_idx, (step_txt, script_txt) in enumerate(act3_steps, start=1):
        c0 = t3.cell(r_idx, 0)
        c1 = t3.cell(r_idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(step_txt)
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(script_txt)
        r1.font.size = Pt(9.5)
        r1.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ACT 4
    add_heading_2("Act 4: High-Risk Policy Violation & Boundary Interception")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("📍 ").bold = True
    p.add_run("Frontend Filter: ").bold = True
    p.add_run("Category: Violation / Event: SECURITY_VIOLATION_DETECTED\n")
    p.add_run("🎯 Objective: ").bold = True
    p.add_run("Demonstrate real-time threat detection, CRITICAL risk badges, and containment.")

    t4 = doc.add_table(rows=2, cols=2)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, text in enumerate(headers):
        cell = t4.cell(0, c_idx)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_h = p_h.add_run(text)
        r_h.bold = True
        r_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h.font.size = Pt(10)

    act4_steps = [
        (
            "Step 4.1: Inspect Policy Violation Event\nFiltering stream for SECURITY_VIOLATION_DETECTED (Agent: agent_sec_auditor).",
            "\"What happens when an agent violates enterprise boundaries? Here in Stream 2, agent_sec_auditor attempted unauthorized PII export. GuardianIQ immediately flagged a CRITICAL risk level, registered a SECURITY_VIOLATION_DETECTED event, and engaged boundary containment—blocking the data transfer before it reached an external endpoint.\""
        )
    ]
    for r_idx, (step_txt, script_txt) in enumerate(act4_steps, start=1):
        c0 = t4.cell(r_idx, 0)
        c1 = t4.cell(r_idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "F8FAFC")
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(step_txt)
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(script_txt)
        r1.font.size = Pt(9.5)
        r1.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ACT 5
    add_heading_2("Act 5: Dead Letter Queue (DLQ) Operational Recovery (/audit/dead-letter)")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("📍 ").bold = True
    p.add_run("Frontend Route: ").bold = True
    p.add_run("/audit/dead-letter (DeadLetterReviewPage.tsx)\n")
    p.add_run("🎯 Objective: ").bold = True
    p.add_run("Demonstrate zero event loss outbox architecture & interactive RetryActionButton.")

    t5 = doc.add_table(rows=3, cols=2)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, text in enumerate(headers):
        cell = t5.cell(0, c_idx)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_h = p_h.add_run(text)
        r_h.bold = True
        r_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h.font.size = Pt(10)

    act5_steps = [
        (
            "Step 5.1: Navigate to Dead Letter Review\nOpening /audit/dead-letter inspector.",
            "\"In distributed enterprise architectures, downstream brokers can experience temporary network outages. GuardianIQ uses a transactional outbox table pattern. Here on the Dead Letter Review page, we can inspect 2 un-delivered message events targeted for downstream SIEM topics.\""
        ),
        (
            "Step 5.2: Click RetryActionButton\nTriggering manual re-queue for failed event.",
            "\"Watch as I click 'Retry Dispatch'. The button transitions to a spinning 'Retrying...' state, invokes our backend re-queue endpoint, and successfully re-dispatches the event with zero data loss!\""
        )
    ]
    for r_idx, (step_txt, script_txt) in enumerate(act5_steps, start=1):
        c0 = t5.cell(r_idx, 0)
        c1 = t5.cell(r_idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(step_txt)
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(script_txt)
        r1.font.size = Pt(9.5)
        r1.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ACT 6
    add_heading_2("Act 6: Compliance Audit Trail Export Engine (/audit/export)")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("📍 ").bold = True
    p.add_run("Frontend Route & Component: ").bold = True
    p.add_run("/audit/export & ExportModal.tsx\n")
    p.add_run("🎯 Objective: ").bold = True
    p.add_run("Demonstrate automated SOC2 / ISO compliance package generation.")

    t6 = doc.add_table(rows=3, cols=2)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, text in enumerate(headers):
        cell = t6.cell(0, c_idx)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p_h = cell.paragraphs[0]
        p_h.paragraph_format.space_after = Pt(0)
        r_h = p_h.add_run(text)
        r_h.bold = True
        r_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_h.font.size = Pt(10)

    act6_steps = [
        (
            "Step 6.1: Open Export Modal\nClicking 'New Export Request' button.",
            "\"Finally, enterprise auditors require exportable evidence packages. Clicking 'New Export Request' opens our shared ExportModal. I will input Subject Entity: 'agents', ID: 'd6a3cb9e-11a8-4004-b82f-a38e33790df0', select JSON format, and enter Reason: 'Q3 SOC2 Compliance Audit'.\""
        ),
        (
            "Step 6.2: Submit & Download Package\nSubmitting form & clicking Download Export Bundle.",
            "\"As I submit, GuardianIQ launches a streaming export task. Within seconds, our compliance export bundle is generated, complete with SHA-256 verification manifests and full event envelope records ready for external audit presentation.\""
        )
    ]
    for r_idx, (step_txt, script_txt) in enumerate(act6_steps, start=1):
        c0 = t6.cell(r_idx, 0)
        c1 = t6.cell(r_idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(step_txt)
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(script_txt)
        r1.font.size = Pt(9.5)
        r1.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Key Takeaways for Stakeholders
    add_heading_1("3. Executive Wrap-Up & Key Value Takeaways")

    takeaways = [
        ("Cryptographic Immutability", "Every event envelope is signed with a SHA-256 hash upon ingestion, guaranteeing tamper-evident audit logs."),
        ("Multi-Dimensional Traceability", "Instant correlation tracking across microservices allows enterprise teams to trace agent decisions back to exact prompt inputs and policy checks."),
        ("Transactional Outbox Resilience", "Zero data loss architecture guarantees that even under network partitioning, audit events are preserved and retryable via the DLQ UI."),
        ("Audit-Ready Compliance", "Instant asynchronous export engine formats audit trails into standardized JSON/CSV packages with full justification logs for SOC2, HIPAA, and ISO 27001 audits.")
    ]

    for title_t, desc_t in takeaways:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(2)
        p_t.paragraph_format.space_after = Pt(4)
        r_bullet = p_t.add_run("✔ ")
        r_bullet.bold = True
        r_bullet.font.color.rgb = RGBColor(0x05, 0x96, 0x69) # Emerald Green
        
        r_head = p_t.add_run(f"{title_t}: ")
        r_head.bold = True
        r_head.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        r_body = p_t.add_run(desc_t)
        r_body.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.save(output_path)
    print(f"Successfully generated Phase 4 Demo Script Word Document: {output_path}")

if __name__ == "__main__":
    out_dir = r"c:\Users\aayus\Desktop\GuardianIQ--1\docs\Phase 4"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "GuardianIQ_Phase4_Live_Demo_Script.docx")
    build_demo_doc(out_path)
