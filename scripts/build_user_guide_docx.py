import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTE", color_hex="2563EB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=120, bottom=120, left=200, right=150)
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{color_hex}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"📌 {title}: ")
    run_t.bold = True
    run_t.font.name = "Segoe UI"
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor.from_string(color_hex)
    
    run_b = p.add_run(text)
    run_b.font.name = "Segoe UI"
    run_b.font.size = Pt(9.5)
    run_b.font.color.rgb = RGBColor(31, 41, 55)
    
    doc.add_paragraph()

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'<w:left w:val="single" w:sz="16" w:space="0" w:color="64748B"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)
    
    doc.add_paragraph()

def build_docx(md_path: Path, out_path: Path):
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    
    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(31, 41, 55)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)
    
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()
    
    in_code_block = False
    code_buffer = []
    
    in_table = False
    table_rows = []
    
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
            
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
            
        # Handle Markdown Tables
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Table row
            cols = [c.strip() for c in line.strip().split('|')[1:-1]]
            # Check if it's separator row |---|---|
            if all(re.match(r'^:?-+:?$', c) for c in cols):
                i += 1
                continue
            
            table_rows.append(cols)
            # Lookahead: if next line is not table row, flush table
            if i + 1 >= n or not (lines[i+1].strip().startswith('|') and lines[i+1].strip().endswith('|')):
                # Flush table
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.autofit = True
                    
                    for r_idx, row_data in enumerate(table_rows):
                        for c_idx in range(num_cols):
                            cell = tbl.cell(r_idx, c_idx)
                            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                            # clean markdown formatting
                            cell_text = cell_text.replace('<br>', '\n').replace('&rarr;', '→').replace('&to;', '→')
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                            cell_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', cell_text)
                            
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            
                            if r_idx == 0:
                                # Header row
                                set_cell_background(cell, "1E3A8A")
                                set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                                run = p.add_run(cell_text)
                                run.bold = True
                                run.font.name = "Segoe UI"
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                # Data row
                                bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                                set_cell_background(cell, bg)
                                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                                run = p.add_run(cell_text)
                                run.font.name = "Segoe UI"
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(31, 41, 55)
                    
                    doc.add_paragraph()
                    table_rows = []
            i += 1
            continue
            
        stripped = line.strip()
        
        # Empty line
        if not stripped:
            i += 1
            continue
            
        # Horizontal Rule
        if stripped == '---':
            i += 1
            continue
            
        # Heading 1
        if stripped.startswith('# '):
            h_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h_text)
            run.bold = True
            run.font.name = 'Segoe UI'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue
            i += 1
            continue
            
        # Heading 2
        if stripped.startswith('## '):
            h_text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h_text)
            run.bold = True
            run.font.name = 'Segoe UI'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(37, 99, 235) # Bright Slate Blue
            i += 1
            continue
            
        # Heading 3
        if stripped.startswith('### '):
            h_text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h_text)
            run.bold = True
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(15, 23, 42) # Slate Dark
            i += 1
            continue
            
        # Heading 4
        if stripped.startswith('#### '):
            h_text = stripped[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h_text)
            run.bold = True
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(51, 65, 85)
            i += 1
            continue

        # Blockquote / Alert
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            # clean markdown formatting
            quote_text = re.sub(r'\*\*(.*?)\*\*', r'\1', quote_text)
            quote_text = re.sub(r'`(.*?)`', r'\1', quote_text)
            add_callout(doc, quote_text, title="GOVERNANCE ADVISORY", color_hex="2563EB")
            i += 1
            continue
            
        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            
            # Simple bold/inline parsing
            parts = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', item_text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(185, 28, 28)
                elif part.startswith('[') and ']' in part:
                    label = re.match(r'\[(.*?)\]', part).group(1)
                    r = p.add_run(label)
                    r.font.color.rgb = RGBColor(37, 99, 235)
                else:
                    p.add_run(part)
            i += 1
            continue
            
        # Numbered list
        m_num = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m_num:
            num_idx, item_text = m_num.groups()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            
            parts = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', item_text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(185, 28, 28)
                elif part.startswith('[') and ']' in part:
                    label = re.match(r'\[(.*?)\]', part).group(1)
                    r = p.add_run(label)
                    r.font.color.rgb = RGBColor(37, 99, 235)
                else:
                    p.add_run(part)
            i += 1
            continue
            
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        
        parts = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', stripped)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.bold = True
            elif part.startswith('`') and part.endswith('`'):
                r = p.add_run(part[1:-1])
                r.font.name = 'Consolas'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(185, 28, 28)
            elif part.startswith('[') and ']' in part:
                label = re.match(r'\[(.*?)\]', part).group(1)
                r = p.add_run(label)
                r.font.color.rgb = RGBColor(37, 99, 235)
            else:
                p.add_run(part)
        
        i += 1
        
    doc.save(out_path)
    print(f"[OK] Generated Word document: {out_path}")

if __name__ == '__main__':
    root = Path(__file__).resolve().parent.parent
    src_md = root / 'USER_GUIDE.md'
    out_docx_root = root / 'GuardianIQ_User_Guide.docx'
    out_docx_docs = root / 'docs' / 'GuardianIQ_User_Guide.docx'
    
    build_docx(src_md, out_docx_root)
    build_docx(src_md, out_docx_docs)

