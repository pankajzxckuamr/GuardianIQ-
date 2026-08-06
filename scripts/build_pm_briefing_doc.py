import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    
    # Left border styling
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="1E3A8A"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    set_cell_margins(cell, top=120, bottom=120, left=200, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"[{title}] ")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10.5)
    run_b.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def build_docx(output_path):
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Executive Briefing: GuardianIQ Workspace Shell Pages Audit")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("Analysis of Static Placeholders, Phase 3/4 Scope Compliance & Phase 5 Roadmap Recommendations")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    # Metadata Block Table
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Target Audience:", "Product Management (PM) & Engineering Leadership"),
        ("Prepared For:", "Product Manager Assistance & Roadmap Alignment"),
        ("System / Repository:", "GuardianIQ Platform (Phases 3 & 4 Audit)"),
        ("Date & Status:", "August 2026 | Verified Complete (Zero Implementation Gaps)")
    ]
    for i, (k, v) in enumerate(meta_data):
        cell_k, cell_v = meta_tbl.cell(i, 0), meta_tbl.cell(i, 1)
        cell_k.width, cell_v.width = Inches(2.0), Inches(4.5)
        set_cell_background(cell_k, "F8FAFC")
        set_cell_background(cell_v, "F8FAFC")
        set_cell_margins(cell_k, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_v, top=60, bottom=60, left=100, right=100)
        
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        rv.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 1: Executive Summary
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r1 = h1.add_run("1. Executive Summary")
    r1.font.name = "Arial"
    r1.font.size = Pt(16)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p = doc.add_paragraph(
        "This briefing document provides a technical audit of the static placeholder pages in GuardianIQ: "
        "the Policies Dashboard (/policies), Risk Workspace (/risk), Approval Hub (/approvals), and Admin Panel (/admin/*). "
        "The purpose of this document is to clarify why these pages currently exist as static UI shells, verify their alignment "
        "against Phase 3 and Phase 4 delivery contracts, and present actionable recommendations for Product Management (PM) planning."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_callout(
        doc,
        "Zero Compliance Gaps Found: All 4 workspace shell pages were deliberately created during Phase 3 as protected routing placeholders. "
        "Phase 3 and Phase 4 deliverables met 100% of their specified acceptance criteria. Full UI mutation workflows for policy building "
        "and custom admin settings were explicitly scheduled in project specs for post-Phase 4 roadmap iterations.",
        title="KEY FINDING FOR PM"
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 2: Detailed Inventory of Workspace Shell Pages
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)
    r2 = h2.add_run("2. Inventory & Current State of Shell Pages")
    r2.font.name = "Arial"
    r2.font.size = Pt(16)
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p = doc.add_paragraph(
        "During Phase 3 implementation, top-level routes were established in AppRouter.tsx under <ProtectedRoute><AppShell> "
        "to ensure clean application navigation without broken routing links (404 errors). The following four routes currently render static layout stubs:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Table of Shell Pages
    tbl_shells = doc.add_table(rows=5, cols=4)
    tbl_shells.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Route URL", "Component File", "Current Implementation", "Technical Status"]
    col_widths = [Inches(1.2), Inches(1.8), Inches(2.3), Inches(1.2)]
    
    # Format Table Header
    hdr_cells = tbl_shells.rows[0].cells
    for j, h_text in enumerate(headers):
        hdr_cells[j].width = col_widths[j]
        set_cell_background(hdr_cells[j], "1E3A8A")
        set_cell_margins(hdr_cells[j], top=100, bottom=100, left=100, right=100)
        p_hdr = hdr_cells[j].paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_hdr = p_hdr.add_run(h_text)
        r_hdr.bold = True
        r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_hdr.font.size = Pt(9.5)
        
    shell_rows = [
        ("/policies", "PoliciesDashboardPage.tsx", "Renders ModuleLayout with static cards ('12 Active Policies', '100% Compliance'). No backend API integration.", "Static Shell"),
        ("/risk", "RiskWorkspacePage.tsx", "Renders ModuleLayout with static metric cards ('LOW Risk Score', '18 Scanned Assets'). No backend API integration.", "Static Shell"),
        ("/approvals", "ApprovalHubPage.tsx", "Renders ModuleLayout displaying static placeholder text ('You have no pending governance approvals').", "Static Shell"),
        ("/admin/*", "AdminPanelPage.tsx", "Renders ModuleLayout with static administrative sub-navigation tabs (Users, Roles, Settings).", "Static Shell")
    ]
    
    for i, row_data in enumerate(shell_rows):
        row_cells = tbl_shells.rows[i+1].cells
        bg_color = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, text in enumerate(row_data):
            row_cells[j].width = col_widths[j]
            set_cell_background(row_cells[j], bg_color)
            set_cell_margins(row_cells[j], top=80, bottom=80, left=100, right=100)
            p_cell = row_cells[j].paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            r_cell = p_cell.add_run(text)
            r_cell.font.size = Pt(9.5)
            if j == 3:
                r_cell.bold = True
                r_cell.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Orange/Amber badge text
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Phase 3 & 4 Verification vs Acceptance Criteria
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(6)
    r3 = h3.add_run("3. Specification & Acceptance Criteria Verification")
    r3.font.name = "Arial"
    r3.font.size = Pt(16)
    r3.bold = True
    r3.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p = doc.add_paragraph(
        "A rigorous audit of the Phase 3 Detailed Task Plan (Phase3Plan_Tasks_Detail.md) and Phase 4 Task Specifications confirms "
        "that these static pages perfectly fulfill their contract terms without missing features."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Bullet points for specs
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1.paragraph_format.space_after = Pt(4)
    r = bp1.add_run("Phase 3 Deliverable Scope (Tasks 1.6 & 2.6): ")
    r.bold = True
    bp1.add_run(
        "The frontend scope was strictly defined to build functional components for the Governance Registry & Relationships "
        "(Relationship Explorer, Add Relationship Wizard, Visual Graph Viewer, Object Responsibility Panel, and Audit Timeline). "
        "Task 2.6 explicitly mandated creating 'top-level routing shells for non-registry workspace routes'."
    )
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2.paragraph_format.space_after = Pt(4)
    r = bp2.add_run("Database Stub Policy Bindings (Task 1.4): ")
    r.bold = True
    bp2.add_run(
        "Task 1.4 defined policy_bindings and evidence_links as 'stub / bridge' tables required in the DB schema for future expansion. "
        "The DDL migrations, SQLAlchemy models, and read-only GET endpoints were built in backend/app/modules/relationship/ as specified."
    )
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3.paragraph_format.space_after = Pt(8)
    r = bp3.add_run("Phase 4 Governance & Telemetry Scope: ")
    r.bold = True
    bp3.add_run(
        "Phase 4 focused on physical event store implementation (governance_events), transactional outbox dispatcher, dead-letter recovery, "
        "audit export package generator, and turning AuditPage.tsx and DashboardPage.tsx into live telemetry tools. Policy frontend builder forms "
        "were documented in phase4_readiness.txt as backlog deferred to post-Phase 4."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 4: Live Functional Counterparts in the Application
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(16)
    h4.paragraph_format.space_after = Pt(6)
    r4 = h4.add_run("4. Live Functional Counterparts in GuardianIQ")
    r4.font.name = "Arial"
    r4.font.size = Pt(16)
    r4.bold = True
    r4.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p = doc.add_paragraph(
        "To avoid confusion, it is important for Product Management to note that while the shell pages listed in Section 2 are static, "
        "their underlying functional capabilities exist and are fully operational across dedicated modules in GuardianIQ:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    # Table of Functional Counterparts
    tbl_func = doc.add_table(rows=5, cols=3)
    tbl_func.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_headers = ["Domain / Capability", "Placeholder Shell Path", "Fully Functional Operational UI Path"]
    f_widths = [Inches(1.8), Inches(1.8), Inches(2.9)]
    
    hdr_cells_f = tbl_func.rows[0].cells
    for j, h_text in enumerate(f_headers):
        hdr_cells_f[j].width = f_widths[j]
        set_cell_background(hdr_cells_f[j], "0F766E") # Teal header
        set_cell_margins(hdr_cells_f[j], top=100, bottom=100, left=100, right=100)
        p_hdr = hdr_cells_f[j].paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_hdr = p_hdr.add_run(h_text)
        r_hdr.bold = True
        r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_hdr.font.size = Pt(9.5)
        
    func_rows = [
        ("Approval Management", "/approvals (ApprovalHubPage)", "/schedule-approvals (ScheduleApprovalQueue.tsx) - Full governance approval workflow queue."),
        ("User & Role Administration", "/admin/users, /admin/roles (AdminPanel)", "/registry/users-roles (RegistryUsersRolesPage.tsx) - Full ABAC user & role management."),
        ("Tenant Administration", "/admin/settings (AdminPanel)", "/tenants (TenantsPage.tsx) - Multi-tenant configuration and tenant isolation management."),
        ("Policy Enforcement", "/policies (PoliciesDashboard)", "Backend Code-level Enforcement: ABAC (abac_service.py), Agent Boundary Checker (boundary_checker.py), Schedule Rules.")
    ]
    
    for i, row_data in enumerate(func_rows):
        row_cells = tbl_func.rows[i+1].cells
        bg_color = "F0FDF4" if i % 2 == 0 else "FFFFFF"
        for j, text in enumerate(row_data):
            row_cells[j].width = f_widths[j]
            set_cell_background(row_cells[j], bg_color)
            set_cell_margins(row_cells[j], top=80, bottom=80, left=100, right=100)
            p_cell = row_cells[j].paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            r_cell = p_cell.add_run(text)
            r_cell.font.size = Pt(9.5)
            if j == 2:
                r_cell.bold = True
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 5: Roadmap Recommendations for Product Management (PM)
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(16)
    h5.paragraph_format.space_after = Pt(6)
    r5 = h5.add_run("5. Recommendations & Phase 5 Roadmap Plan")
    r5.font.name = "Arial"
    r5.font.size = Pt(16)
    r5.bold = True
    r5.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p = doc.add_paragraph(
        "For upcoming Phase 5 sprint planning, Product Management can prioritize turning these shell pages into full UI features "
        "by leveraging the backend infrastructure already completed in Phase 4. Below is the proposed feature implementation plan:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    rec_items = [
        ("Policy Management Module (/policies)", 
         "Build an interactive Visual Policy Builder form UI. Connect frontend to backend/app/modules/relationship/ models for creating, editing, and binding policy_bindings and evidence_links directly to registry entities."),
        ("Consolidated Approval Hub (/approvals)", 
         "Expand the current /schedule-approvals queue into a unified multi-domain Approval Hub (/approvals) that handles schedule requests, boundary override approvals, and policy exception requests under one view."),
        ("Risk Assessment Workspace (/risk)", 
         "Connect /risk to the backend risk score calculation engine and ABAC clearance metrics to render automated security risk heatmaps and asset risk ratings."),
        ("Unified Admin Settings (/admin/*)", 
         "Consolidate tenant management (/tenants) and RBAC user management (/registry/users-roles) into the tabbed sub-navigation of /admin for a seamless enterprise administration experience.")
    ]
    
    for title, desc in rec_items:
        p_rec = doc.add_paragraph(style='List Bullet')
        p_rec.paragraph_format.space_after = Pt(6)
        r_t = p_rec.add_run(f"{title}: ")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        p_rec.add_run(desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    
    # Save Document
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    output_docx_path = r"c:\Users\aayus\Desktop\GuardianIQ--1\GuardianIQ_Workspace_Shell_Pages_PM_Briefing.docx"
    build_docx(output_docx_path)
