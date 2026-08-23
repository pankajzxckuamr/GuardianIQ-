from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = Path('docs/sip_final_report.md')
out = Path('docs/sip_final_report.docx')

text = src.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for line in lines:
    if not line.strip():
        doc.add_paragraph()
        continue
    if line.startswith('# '):
        p = doc.add_paragraph()
        p.add_run(line[2:]).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(line[3:])
        run.bold = True
        run.font.size = Pt(13)
    elif line.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(line[4:])
        run.bold = True
        run.font.size = Pt(12)
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:])
    elif line.startswith(tuple(f'{i}. ' for i in range(1, 10))):
        p = doc.add_paragraph(style='List Number')
        p.add_run(line)
    elif line.startswith('---'):
        doc.add_paragraph()
    else:
        doc.add_paragraph(line)

doc.save(out)
print(out)
