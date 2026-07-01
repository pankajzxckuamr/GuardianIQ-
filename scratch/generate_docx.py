import os
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding/margins in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("GuardianIQ — Phase 2 Completion")
    title_run.font.name = 'Calibri Light'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Project Learnings and Material Allocation Guide")
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add a horizontal rule
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="1F4E78"/>'
                     r'</w:pBdr>')
    pPr.append(pBdr)

    # Instructions Callout
    callout_table = doc.add_table(rows=1, cols=1)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout_table.autofit = False
    
    callout_cell = callout_table.rows[0].cells[0]
    callout_cell.width = Inches(6.5)
    set_cell_background(callout_cell, "F2F5F8")
    set_cell_margins(callout_cell, top=140, bottom=140, left=200, right=200)
    
    # Left border for callout
    tcPr = callout_cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                          r'<w:left w:val="single" w:sz="24" w:space="0" w:color="007ACC"/>'
                          r'<w:top w:val="none"/>'
                          r'<w:bottom w:val="none"/>'
                          r'<w:right w:val="none"/>'
                          r'</w:tcBorders>')
    tcPr.append(tcBorders)
    
    cp = callout_cell.paragraphs[0]
    c_run1 = cp.add_run("Instructions for use: ")
    c_run1.bold = True
    c_run1.font.color.rgb = RGBColor(0x00, 0x7A, 0xCC)
    c_run2 = cp.add_run("This document outlines the learnings from the GuardianIQ Phase 2 completion data, organized exactly to match your Excel spreadsheet columns. Use Table 1 below to copy and paste the rows directly into your tracking sheet, or use the descriptive guides below to understand the structure.")
    c_run2.italic = True

    doc.add_paragraph() # Spacing

    # Section Heading
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("Table 1: Learnings Matrix (For Spreadsheet Import)")
    h1_run.font.name = 'Calibri Light'
    h1_run.font.size = Pt(15)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # Matrix Table Data
    data = [
        ("1", "WBS 2.3 & 2.5 (Workflow Lifecycle & State Machines)", 
         "Implementation of complex scheduling and execution workflows without detailed specifications leads to gaps in status transition validation, timezone handling, and execution sequencing.", 
         "Conducted a thorough audit of the scheduler and run services against the Phase 2 Developer Spec, adding rigorous validations for cron patterns, timezone settings, and state machine transitions.", 
         "Understanding the implementation and workflow is more important before coding. Specifications act as a blueprint to avoid architectural misalignment and logic gaps.", 
         "AI Developer"),
        ("2", "WBS 2.7 (RBAC & ABAC Authorization)", 
         "Securing sensitive API endpoints and restricting action buttons in UI (add, edit, delete, run-now) based on roles, hierarchy, and context without a standardized access control model.", 
         "Implemented a comprehensive Authorization Decision Service combining RBAC (role-based) and ABAC (attribute-based, department, risk levels, and delegations) checks. Added screen-level button visibility controls and an authorization simulator.", 
         "Hierarchy, add, edit, and ABAC control mechanisms must map to a screen for each user type. Restricting views and actions based on user contexts prevents unauthorized overrides.", 
         "AI Developer"),
        ("3", "WBS 2.1 (Database Schema and Migration)", 
         "Database tables lacked constraints, proper indexes on high-query JSONB columns, and automatic timestamp updates, risking data corruption and poor performance.", 
         "Added schema migrations defining tables for schedules, assignments, runs, steps, and audit events. Configured unique keys, GIN indexes on JSONB fields, and updated_at triggers.", 
         "Database schema integrity is the foundation of AI governance. Applying constraints, GIN indexes on JSONB, and automated audit triggers ensures reliable execution tracking.", 
         "AI Developer"),
        ("4", "WBS 2.2 (SQLAlchemy Models and validation)", 
         "Risk of saving invalid cron schedules, mismatched execution modes, or incorrect timezone configurations due to lack of validation on input payloads.", 
         "Implemented SQLAlchemy models using Postgres-specific JSONB fields and designed Pydantic request schemas with custom validators for cron syntax, timezone verification, and confidence boundaries.", 
         "Moving validation to the schema input boundary (Pydantic field validators) catches invalid configurations before they hit the database, reducing backend complexity.", 
         "AI Developer"),
        ("5", "WBS 2.4 (AI Agent Boundary Checker)", 
         "AI agents could execute unauthorized tools or exceed their permitted execution modes, leading to security breaches (e.g. executing write/execute tools).", 
         "Created the BoundaryChecker service to statically and dynamically validate tool calls, block execution if requested mode exceeds agent limits, and flag schedules containing write-capable tools for human approval.", 
         "Secure AI automation requires active boundary checks at both assignment (configuration) and execution (runtime) stages. Any write-capable tools must automatically trigger human-in-the-loop approvals.", 
         "AI Developer"),
        ("6", "WBS 2.5 (Run Engine and Steps Timeline)", 
         "Monitoring complex multi-agent execution was difficult due to lack of visibility into intermediate states, execution times, and transient failure recovery.", 
         "Built a 6-step run engine generating 4 distinct audit steps (POLICY_CHECK, BOUNDARY_CHECK, AGENT_INVOCATION, OUTPUT_PARSING) persisted to the database. Added SLA breach tracking and retry policies.", 
         "Breaking down agent execution into discrete, persisted run steps provides deep visibility. Tracking failures by types (validation, SLA, boundary, database) allows fine-grained error recovery.", 
         "AI Developer"),
        ("7", "WBS 2.6 (Scheduler Worker Concurrency)", 
         "Running multiple worker instances in a horizontal scaling setup caused duplicate runs of the same due schedule.", 
         "Rewrote the due schedule query to use database-level locking FOR UPDATE SKIP LOCKED and updated next_run_at inside the locked transaction.", 
         "In a distributed deployment, atomic locking (FOR UPDATE SKIP LOCKED) is mandatory to guarantee exact-once execution and prevent duplicate triggers.", 
         "AI Developer"),
        ("8", "WBS 2.8 (Centralized Governance Event Logs)", 
         "Lack of a centralized, tampering-resistant audit log for actions like schedule creation, activation, run failure, and boundary checks.", 
         "Developed the GovernanceEventService to publish 19 distinct event codes to a unified audit_events table and attached a database trigger that rejects updates or deletes (immutability).", 
         "Compliance logs must be immutable. Implementing database-level triggers that block updates/deletes ensures audit trail integrity, satisfying stringent security regulations.", 
         "AI Developer"),
        ("9", "WBS 2.9 & 2.10 (Create Wizard & Dashboard UX)", 
         "Configuring schedules, agent assignments, and boundaries is complex for end-users, leading to setup mistakes.", 
         "Designed a 6-step Create Schedule Wizard including a CronExpressionBuilder, dynamic risk-warning banners, and permission-based action visibility on the dashboard.", 
         "Intricate security policies (like tool blocks and risk levels) must be guided by clean wizard-based UI designs, alert banners, and interactive builders to eliminate user configuration errors.", 
         "AI Developer"),
        ("10", "WBS 2.11 & 2.12 (Run Detail & Approval Queue UI)", 
         "Approvers and auditors lacked intuitive tools to review pending changes, see run steps timeline, or inspect evidence.", 
         "Built the Approval Queue table, details panel (highlighting write tools and risk warnings), and Run Detail Page showing an interactive steps timeline, findings, and JSON evidence viewer.", 
         "Visualizing the execution timeline (with colored status states) and highlighting high-risk actions (in amber/red) makes human review efficient and reduces sign-off delays.", 
         "AI Developer"),
        ("11", "WBS 2.13 (Test Suite Integrity)", 
         "Hidden setup bugs in the testing suite (like the tool allowlist default override) caused false positives, indicating security checks were working when they were bypassed.", 
         "Patched tool allowlist logic in test fixtures to handle empty lists explicitly and expanded authorization negative tests to cover 22 scenarios.", 
         "Test suites must be audited as closely as production code. Default overrides in test fixtures can lead to silent test coverage gaps, especially for security assertions.", 
         "AI Developer"),
        ("12", "WBS 2.14 & 2.15 (Deployment and Dev Documentation)", 
         "Onboarding new developers and deploying to production was prone to human error due to the lack of clear setup steps and architectural maps.", 
         "Documented a detailed 13-step startup sequence in DEPLOYMENT.md, drew Mermaid architecture diagrams, and added interactive help guides (ScreenGuides) across all frontend screens.", 
         "A project is not complete without comprehensive developer and deployment documentation. Contextual help widgets in the UI bridge the gap between technical specs and user operations.", 
         "AI Developer")
    ]

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ["S. No.", "Reference", "Problem Statement", "Solution Provided", "Learning", "Updated By"]
    hdr_cells = table.rows[0].cells
    col_widths = [Inches(0.5), Inches(1.1), Inches(1.4), Inches(1.4), Inches(1.4), Inches(0.7)]
    
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "1F4E78")
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
        hdr_cells[i].width = col_widths[i]

    # Content rows
    for idx, row_data in enumerate(data):
        row = table.add_row()
        cells = row.cells
        is_even = (idx % 2 == 1)
        bg_color = "F2F5F8" if is_even else "FFFFFF"
        
        for i, val in enumerate(row_data):
            cells[i].text = val
            p = cells[i].paragraphs[0]
            p.style = doc.styles['Normal']
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(8.5)
            set_cell_background(cells[i], bg_color)
            set_cell_margins(cells[i], top=80, bottom=80, left=100, right=100)
            cells[i].width = col_widths[i]

    doc.add_paragraph() # Spacing

    # Section 2: Allocation Guide
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Section 2: Material Allocation Guide (Column Breakdown)")
    h2_run.font.name = 'Calibri Light'
    h2_run.font.size = Pt(15)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    columns_desc = [
        ("Column 1: S. No.", "Enter chronological numbering starting from 1 up to 12."),
        ("Column 2: Reference", "Put the module or design document package reference here. This maps the learnings directly back to the project structure (e.g., database schema, scheduling engines, API design, security frameworks, frontend UI, or DevOps configuration)."),
        ("Column 3: Problem Statement", "Put the technical challenge or functional requirement gap that the project had to resolve during Phase 2. This describes why the task was necessary and what would fail if it weren't implemented."),
        ("Column 4: Solution Provided", "Put the engineering solution that was implemented in the code (e.g., adding triggers, building services, creating new UI components, or rewriting SQL queries to use locks) to address the corresponding Problem Statement."),
        ("Column 5: Learning", "Put the development takeaway, security rule, or architectural best practice discovered by the team while implementing the solution. Note: Rows 1 and 2 incorporate and expand upon your handwritten draft notes (\"Understanding the implementation and workflow...\" and \"Heirarchy, add, edit, abac...\") into clear, professional statements."),
        ("Column 6: Updated By", "Insert the name of the role or team member who executed the work and documented the learnings (e.g., \"AI Developer\", \"Aayush\", or the designated QA Lead).")
    ]

    for title, desc in columns_desc:
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(title)
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p_title.paragraph_format.space_before = Pt(4)
        p_title.paragraph_format.space_after = Pt(2)
        
        p_desc = doc.add_paragraph(desc)
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after = Pt(6)

    # Save the file
    out_dir = r"c:\Users\aayus\Desktop\GuardianIQ--1\docs"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Learnings_Phase2_Completion.docx")
    doc.save(out_path)
    print(f"Successfully generated DOCX at {out_path}")

if __name__ == "__main__":
    create_document()
